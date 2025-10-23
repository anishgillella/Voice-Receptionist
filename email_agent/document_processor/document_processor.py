"""Document processor for extracting text from various file formats."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processor for extracting text from documents."""

    SUPPORTED_EXTENSIONS = ["pdf", "docx", "doc"]

    @staticmethod
    def get_file_extension(filename: str) -> Optional[str]:
        """Get file extension."""
        ext = Path(filename).suffix.lower().lstrip(".")
        if ext in DocumentProcessor.SUPPORTED_EXTENSIONS:
            return ext
        return None

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            from io import BytesIO

            pdf_reader = PdfReader(BytesIO(file_content))
            text = ""

            for page in pdf_reader.pages:
                text += page.extract_text()

            logger.info(f"Successfully extracted text from PDF ({len(text)} characters)")
            return text
        except Exception as error:
            logger.error(f"Error extracting text from PDF: {error}")
            return None

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            from io import BytesIO

            doc = Document(BytesIO(file_content))
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"

            logger.info(f"Successfully extracted text from DOCX ({len(text)} characters)")
            return text
        except Exception as error:
            logger.error(f"Error extracting text from DOCX: {error}")
            return None

    @staticmethod
    def extract_text_from_doc(file_content: bytes) -> Optional[str]:
        """Extract text from DOC file (legacy Word format)."""
        try:
            # DOC files (older Word format) are harder to parse
            # We'll attempt to use python-docx which may handle some .doc files
            from io import BytesIO

            doc = Document(BytesIO(file_content))
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"

            logger.info(f"Successfully extracted text from DOC ({len(text)} characters)")
            return text
        except Exception as error:
            logger.error(f"Error extracting text from DOC: {error}")
            # Note: For older .doc files, you may need to use python-docx2docx or similar
            return None

    @staticmethod
    def extract_text(filename: str, file_content: bytes) -> Optional[str]:
        """
        Extract text from a document based on file extension.

        Args:
            filename: Name of the file
            file_content: Binary content of the file

        Returns:
            Extracted text or None if extraction failed
        """
        extension = DocumentProcessor.get_file_extension(filename)

        if extension == "pdf":
            return DocumentProcessor.extract_text_from_pdf(file_content)
        elif extension == "docx":
            return DocumentProcessor.extract_text_from_docx(file_content)
        elif extension == "doc":
            return DocumentProcessor.extract_text_from_doc(file_content)
        else:
            logger.warning(f"Unsupported file format: {filename}")
            return None

    @staticmethod
    def get_file_metadata(filename: str, file_size_bytes: int) -> dict:
        """Get metadata about a document."""
        extension = DocumentProcessor.get_file_extension(filename)

        return {
            "filename": filename,
            "extension": extension,
            "file_size_mb": round(file_size_bytes / (1024 * 1024), 2),
            "file_size_bytes": file_size_bytes,
            "is_supported": extension is not None,
        }
